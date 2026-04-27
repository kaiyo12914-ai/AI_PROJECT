import os
import django
from django.conf import settings
from docx import Document as DocxDocument
import io

# 設置 Django 的環境變數，如果尚未配置過
if not settings.configured:
    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "webproj.test_settings_final")

# 初始化 Django
django.setup()

from webapps.comment.domain.services.evaluation_service import EvaluationService
from webapps.comment.domain.entities.evaluation import Evaluation

def test_generate_consideration_sheet():
    """測試 generate_consideration_sheet_from_template 是否能讀取樣版檔案並產生輸出"""
    template_path = os.path.join(
        os.getcwd(),
        r"webapps\comment\data\templates_office\考績表_template.docx"
    )

    # 檢查樣版文件是否存在
    if not os.path.exists(template_path):
        print(f"模板檔案不存在：{template_path}")
        return

    try:
        evaluations = list(Evaluation.objects.order_by('-created_at'))

        service = EvaluationService()
        result_doc = service.generate_consideration_sheet_from_template(
            template_path=template_path,
            evaluations=evaluations
        )
        
        result_docx = DocxDocument(io.BytesIO(result_doc))
        result_docx.save("template_output.docx")
        print(f"成功生成考績表，大小：{len(result_doc)} bytes")

    except Exception as e:
        print(f"測試失敗：{str(e)}")

def test_add_coordinates_to_cells():
    # 使用範例
    doc_path = os.path.join(
        os.getcwd(),
        r"webapps\comment\data\templates_office\考績表_template.docx"
    )
    output_path = 'output.docx'
    
    # 打開文檔
    doc = DocxDocument(doc_path)

    # 遍歷文檔中的所有表格
    for table in doc.tables:
        # 遍歷表格中的所有行
        for row_idx, row in enumerate(table.rows):
            # 遍歷行中的所有單元格
            for col_idx, cell in enumerate(row.cells):
                # 獲取單元格中的現有文字
                original_text = cell.text
                # 將座標值添加到單元格中
                new_text = f"({row_idx}, {col_idx}) {original_text}"
                cell.text = new_text

    # 保存修改後的文檔
    doc.save(output_path)
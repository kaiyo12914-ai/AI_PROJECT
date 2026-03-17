import docx
import sys
import os

# Set console output to UTF-8
if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

try:
    import glob
    files = glob.glob(r'C:\Users\kaiyo12914\Downloads\*.docx')
    target = ""
    for f in files:
        # Check by name manually or use the last modified one
        if "~оǤզW.docx" in f or "оǤ" in f:
            target = f
    
    if not target:
        # Fallback to any docx in that folder for testing
        docx_files = [f for f in files if f.endswith('.docx')]
        if docx_files:
            target = docx_files[0]

    if not target:
        print("Error: No docx found.")
        sys.exit(1)
        
    doc = docx.Document(target)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        full_text.append("--- TABLE START ---")
        for row in table.rows:
            # Deduplicate text if cells are merged
            row_data = []
            last_text = None
            for cell in row.cells:
                text = cell.text.strip()
                if text != last_text:
                    row_data.append(text)
                    last_text = text
            full_text.append(" | ".join(row_data))
            
    print("\n".join(full_text))
except Exception as e:
    print(f"Error: {str(e)}")

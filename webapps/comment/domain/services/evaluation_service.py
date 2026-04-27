import os
import re
import json
from typing import List
from django.conf import settings

from docx import Document as DocxDocument
from docx.table import Table,_Cell
from docx.document import Document as DocumentObject
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from webapps.comment.domain.entities.evaluation import Evaluation
from webapps.comment.domain.model.SubPerformanceGrades import SubPerformanceGrades

class EvaluationService:
    """服務類別：處理評語相關業務邏輯"""

    def __init__(self) -> None:
        
        # 身份證字號的單元格位置    
        self.identity_cell_row = 10
        self.identity_cell_col = 9

        # 評語格位置
        self.comment_row_idx = 27
        self.comment_col_idx = 31
    
    def fetch_evaluations_for_user(self,creator_account: str) -> List[Evaluation]:
        """
        取得當前使用者所建立的所有評語。

        Args:
            creator_account (str): 使用者帳號

        Returns:
            List[Evaluation]: 評語列表
        """
        from webapps.comment.domain.entities.evaluation import Evaluation
        return list(Evaluation.objects.filter(creator_account=creator_account)
                                  .order_by('-created_at'))
    
    def generate_consideration_sheet_from_template(self,template_path: str, evaluations: List[Evaluation]) -> bytes:
        """
        根據樣版檔案及評語列表，生成考績表Excel/Word檔。
        本功能會在每個頁面中尋找身份證字號並填入對應的評語資料。

        Args:
            template_path (str): 樣版檔案路徑
            evaluations (List[Evaluation]): 評語資料

        Returns:
            bytes: 生成後的二進位檔案內容

        Note:
            暫時實作為Word格式，未來可考慮支援多種輸出格式。
        """

        try:
            # 確認樣版檔存在
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")

            # 載入樣版檔案
            template_doc = DocxDocument(template_path)            

            # 檢查樣版是否有表格
            if not template_doc.tables:
                raise ValueError(f"Template has no tables: {template_path}")

            # 遍歷樣版中的所有頁面（每個表格為一頁）
            self.ProcessAllTables(evaluations, template_doc)

            import io
            buffer = io.BytesIO()
            template_doc.save(buffer)            
            docx_bytes = buffer.getvalue()            
            buffer.close()

            return docx_bytes

        except Exception as e:
            raise RuntimeError(f"Failed to generate consideration sheet: {str(e)}")
    
    def ProcessAllTables(self,evaluations:List[Evaluation], template_doc:DocumentObject):        
        for table in template_doc.tables:            

            # 儀態
            cell_bodyTypeModerate = table.cell(12,15)
            cell_bodyTypeSlim = table.cell(12,19)
            cell_bodyTypeOverweight = table.cell(12,22)
            cell_bodyTypeUnderweight = table.cell(12,26)
            cell_bodyTypeObese = table.cell(12,29)
            # 體能
            cell_fitnessQualified = table.cell(13,21)
            cell_fitnessUnqualified = table.cell(13,31)

            subgrade_cells = [
                cell_bodyTypeModerate,
                cell_bodyTypeSlim,
                cell_bodyTypeOverweight,
                cell_bodyTypeUnderweight,
                cell_bodyTypeObese,
                cell_fitnessQualified,
                cell_fitnessUnqualified
            ]

            # 初始化儲存格
            self.InitSubGrades(table,subgrade_cells)

            # 儲存格內容為「身份證字號\n姓名」，需要分割元素後取得身份證字號
            idno_cell_text = table.cell(self.identity_cell_row, self.identity_cell_col).text.strip()
            
            if not re.search(r'[A-Za-z]{1}[0-9]{9}', idno_cell_text):    
                continue

            elements = idno_cell_text.split("\n")
            idno = elements[1]                    

            filtered_evaluations = list(filter(lambda eval:eval.idno == idno ,evaluations))                    
            if len(filtered_evaluations) == 0:
                break
            
            # 找到evaluation後開始填值
            evaluation = filtered_evaluations[0]
            self.ProcessSingleTable(table,evaluation)

            # 勾選考核維度
            subPerformanceGrades = SubPerformanceGrades(
                thoughtGradeOptions=evaluation.subPerformanceGrades[0]['thoughtGradeOptions'],
                moralityOptions=evaluation.subPerformanceGrades[1]['moralityOptions'],
                abilityOptions=evaluation.subPerformanceGrades[2]['abilityOptions'],
                knowledgeOptions=evaluation.subPerformanceGrades[3]['knowledgeOptions'],
                performanceOptions=evaluation.subPerformanceGrades[4]['performanceOptions'],
                postureOptions=evaluation.subPerformanceGrades[5]['postureOptions'],
                physicalOptions=evaluation.subPerformanceGrades[6]['physicalOptions'],
                hrSuggestionOptions=evaluation.subPerformanceGrades[7]['hrSuggestionOptions']
            )

            score_col = {"特優":14,"優等":15,"甲上":17,"甲等":19,"乙上":20,"乙等":22,"丙上":24,"丙等":26,"丁等":27}            
            table.cell(7, score_col[subPerformanceGrades.thoughtGradeOptions]).text = "V"            
            table.cell(8, score_col[subPerformanceGrades.moralityOptions]).text = "V"
            table.cell(9, score_col[subPerformanceGrades.abilityOptions]).text = "V"
            table.cell(10, score_col[subPerformanceGrades.knowledgeOptions]).text = "V"
            table.cell(11, score_col[subPerformanceGrades.performanceOptions]).text = "V"
            
            # 儀態
            postureScores = {"適中":15,"瘦":19,"胖":22,"過瘦":26,"過胖":29}
            table.cell(12, postureScores[subPerformanceGrades.postureOptions]).text = "V"

            # 體能
            physicalScores = {"合格":17,"不合格":27}
            table.cell(13, physicalScores[subPerformanceGrades.physicalOptions]).text = "V"

            # 人事建議(初考官與覆考官都一起設置相同的值，原則上不會有差異)            
            hrScores = {"部隊指揮職":15,"一般幕僚職":17,"專才專業職":20,"向上派職":22,"進修深造":24,"平衡歷練":27,"續任現職":29,"不適現職":30}
            table.cell(18, hrScores[subPerformanceGrades.hrSuggestionOptions]).text = "V"
            table.cell(19, hrScores[subPerformanceGrades.hrSuggestionOptions]).text = "V"

    def ProcessSingleTable(self,table:Table,evaluation:Evaluation):
        # 填入評語
        comment_cell = table.cell(self.comment_row_idx,self.comment_col_idx)                                
        origin_text = comment_cell.text                
        comment_cell.text = ""
        
        run = comment_cell.paragraphs[0].add_run(origin_text + evaluation.comment_text)                 
        # 設定英數為"標楷體"
        run.font.name = "標楷體"
        # 設定中文為"標楷體"
        run._element.rPr.rFonts.set(qn('w:eastAsia'),"標楷體")                   
              
    def InitSubGrades(self,table:Table,subgrade_cells:List[_Cell]):
        """
        清空已打勾的部份
        思想 至 績效 用一個迴圈 處理
        儀態、體能另一個迴圈 處理
        """
        for row_index in range(7,11+1):
            for column_index in range(14,31+1):
                sub_cell = table.cell(row_index,column_index)
                sub_cell.text = ""

        for cell in subgrade_cells:
            cell.text = ""
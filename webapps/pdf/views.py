from __future__ import annotations

import io
import json
import os
from datetime import datetime
from typing import List, Tuple

from django.http import HttpResponse, JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt

from webapps.llm.llm_factory import get_chat_model
from webapps.portal.decorators import require_node


def _norm_base(path: str) -> str:
    s = (path or "").strip()
    if not s:
        return ""
    if not s.startswith("/"):
        s = "/" + s
    while len(s) > 1 and s.endswith("/"):
        s = s[:-1]
    return "" if s == "/" else s


def _calc_app_base_url(request) -> str:
    """
    Compute app base url for apiurl():
    - direct: /pdf/... -> /pdf
    - proxied: /djangoai/pdf/... -> /djangoai/pdf
    """
    script = _norm_base(getattr(request, "script_name", "") or request.META.get("SCRIPT_NAME", ""))
    return _norm_base((script + "/pdf").replace("//", "/"))


@require_node("pdf")
def index(request):
    return render(request, "pdf/index.html", {"app_base_url": _calc_app_base_url(request)})


def _resolve_tesseract_cmd() -> str:
    env_cmd = (os.environ.get("TESSERACT_CMD") or "").strip()
    if env_cmd and os.path.exists(env_cmd):
        return env_cmd

    default_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(default_cmd):
        return default_cmd
    return ""


def _resolve_tessdata_dir(tesseract_cmd: str) -> str:
    env_dir = (os.environ.get("TESSDATA_PREFIX") or "").strip()
    candidates: List[str] = []
    if env_dir:
        candidates.append(env_dir)
        if not env_dir.lower().endswith("tessdata"):
            candidates.append(os.path.join(env_dir, "tessdata"))

    if tesseract_cmd:
        candidates.append(os.path.join(os.path.dirname(tesseract_cmd), "tessdata"))

    candidates.append(r"C:\Program Files\Tesseract-OCR\tessdata")

    for d in candidates:
        if d and os.path.isdir(d):
            return d
    return ""


def _ocr_config(psm: int, tessdata_dir: str) -> str:
    base = f"--oem 3 --psm {psm}"
    if tessdata_dir:
        return f'{base} --tessdata-dir "{tessdata_dir}"'
    return base


def _safe_filename_base(name: str) -> str:
    base = os.path.splitext(os.path.basename(name or "document"))[0]
    for ch in ['"', "'", "\\", "/", ":", "*", "?", "<", ">", "|", "\r", "\n", "\t"]:
        base = base.replace(ch, "_")
    return base or "document"


def _get_uploaded_pdf(request):
    f = request.FILES.get("pdf")
    if not f:
        return None, JsonResponse({"ok": False, "error": "請先上傳 PDF"}, status=400)
    if not f.name.lower().endswith(".pdf"):
        return None, JsonResponse({"ok": False, "error": "僅支援 .pdf 檔案"}, status=400)
    max_mb = int(os.environ.get("PDF_MAX_MB", "50") or 50)
    if f.size > max_mb * 1024 * 1024:
        return None, JsonResponse({"ok": False, "error": f"檔案大小超過上限 {max_mb}MB"}, status=400)
    return f, None


def _extract_pdf_text_pypdf(file_obj) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_obj)
    parts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(f"[第 {i} 頁]\n{t}")
    return "\n\n".join(parts).strip()


def _ocr_pdf_text(file_bytes: bytes, lang: str = "chi_tra+eng") -> str:
    import pypdfium2 as pdfium
    import pytesseract

    tcmd = _resolve_tesseract_cmd()
    if not tcmd:
        raise RuntimeError("找不到 tesseract.exe，請設定 TESSERACT_CMD 或確認安裝路徑")
    pytesseract.pytesseract.tesseract_cmd = tcmd

    tessdata_dir = _resolve_tessdata_dir(tcmd)
    if not tessdata_dir:
        raise RuntimeError("找不到 tessdata 目錄，請設定 TESSDATA_PREFIX 指向 tessdata")
    os.environ["TESSDATA_PREFIX"] = tessdata_dir

    pdf = pdfium.PdfDocument(file_bytes)
    parts: List[str] = []
    for i in range(len(pdf)):
        page = pdf.get_page(i)
        pil_img = page.render(scale=2.5).to_pil()
        text = (
            pytesseract.image_to_string(
                pil_img,
                lang=lang,
                config=_ocr_config(6, tessdata_dir),
            )
            or ""
        ).strip()
        if text:
            parts.append(f"[第 {i + 1} 頁 OCR]\n{text}")
        else:
            parts.append(f"[第 {i + 1} 頁 OCR]\n（未擷取到文字）")
        page.close()
    return "\n\n".join(parts).strip()


def _extract_pdf_text_auto(file_obj, ocr_lang: str = "chi_tra+eng", min_chars: int = 40) -> Tuple[str, bool]:
    file_bytes = file_obj.read()

    text = ""
    try:
        text = _extract_pdf_text_pypdf(io.BytesIO(file_bytes))
    except Exception:
        text = ""

    if text and len(text) >= min_chars:
        return text, False

    try:
        ocr_text = _ocr_pdf_text(file_bytes, lang=ocr_lang)
        return (ocr_text or "").strip(), True
    except Exception as e:
        if text and text.strip():
            return text.strip(), False
        raise RuntimeError(f"PDF OCR 失敗：{e}")


@csrf_exempt
@require_node("pdf", api=True)
def api_extract_text(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)

    f, err = _get_uploaded_pdf(request)
    if err:
        return err

    ocr_lang = (request.POST.get("ocr_lang") or os.environ.get("OCR_LANG") or "chi_tra+eng").strip()
    try:
        text, used_ocr = _extract_pdf_text_auto(f, ocr_lang=ocr_lang)
        return JsonResponse(
            {"ok": True, "filename": f.name, "chars": len(text), "text": text, "used_ocr": used_ocr},
            status=200,
        )
    except Exception as e:
        return JsonResponse({"ok": False, "error": str(e)}, status=500)


@csrf_exempt
@require_node("pdf", api=True)
def api_download_txt(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)

    f, err = _get_uploaded_pdf(request)
    if err:
        return err

    ocr_lang = (request.POST.get("ocr_lang") or os.environ.get("OCR_LANG") or "chi_tra+eng").strip()
    try:
        text, used_ocr = _extract_pdf_text_auto(f, ocr_lang=ocr_lang)
        base = _safe_filename_base(f.name)
        stamp = datetime.now().strftime("%Y%m%d")
        filename = f"{base}_{stamp}.txt"
        header = f"[OCR 模式] lang={ocr_lang}\n\n" if used_ocr else ""
        resp = HttpResponse(header + (text or ""), content_type="text/plain; charset=utf-8")
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp
    except Exception as e:
        return JsonResponse({"ok": False, "error": str(e)}, status=500)


@csrf_exempt
@require_node("pdf", api=True)
def api_download_docx(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)

    f, err = _get_uploaded_pdf(request)
    if err:
        return err

    ocr_lang = (request.POST.get("ocr_lang") or os.environ.get("OCR_LANG") or "chi_tra+eng").strip()
    try:
        from docx import Document

        text, used_ocr = _extract_pdf_text_auto(f, ocr_lang=ocr_lang)

        doc = Document()
        doc.add_heading("PDF 擷取結果", level=1)
        if used_ocr:
            doc.add_paragraph(f"OCR 模式：lang={ocr_lang}")

        if text:
            for block in text.split("\n\n"):
                t = (block or "").strip()
                if t:
                    doc.add_paragraph(t)
        else:
            doc.add_paragraph("未擷取到文字")

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        base = _safe_filename_base(f.name)
        stamp = datetime.now().strftime("%Y%m%d")
        filename = f"{base}_{stamp}.docx"
        resp = HttpResponse(
            bio.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp
    except Exception as e:
        return JsonResponse({"ok": False, "error": str(e)}, status=500)


def _build_summary_prompt(doc_text: str) -> str:
    return f"""你是一位專業的文件分析與摘要專家，請根據以下文件內容，產出結構化重點摘要。

【任務要求】
1. 提取文件的核心主題與目的
2. 梳理關鍵重點（條列式）
3. 保留重要數據、結論與專有名詞
4. 避免冗長敘述，不要逐句改寫
5. 使用清晰、簡潔、專業的語言

【輸出格式】
請嚴格依照以下結構輸出：

一、文件主題：
（用一句話說明）

二、核心摘要：
（3～5句重點總結）

三、關鍵重點：
- 重點1
- 重點2
- 重點3
（視內容增加）

四、重要數據 / 結論：
- （列出關鍵數據或結論）

五、可行動建議（若適用）：
- （可執行的建議）

【文件內容】
{doc_text}
""".strip()


def _llm_to_text(resp) -> str:
    if resp is None:
        return ""
    if hasattr(resp, "content"):
        c = resp.content
        if isinstance(c, list):
            out = []
            for part in c:
                if isinstance(part, str):
                    out.append(part)
                elif isinstance(part, dict):
                    out.append(str(part.get("text") or part.get("content") or ""))
            return "".join(out).strip()
        return str(c or "").strip()
    return str(resp).strip()


def _local_summary_fallback(doc_text: str) -> str:
    lines = [ln.strip() for ln in doc_text.splitlines() if ln.strip()]
    topic = lines[0] if lines else "未能判定主題"
    key_points = lines[:6]
    top_data = [ln for ln in lines if any(ch.isdigit() for ch in ln)][:5]

    summary_lines = []
    if lines:
        summary_lines.append(f"文件主要圍繞「{topic}」展開。")
        summary_lines.append("內容已完成初步結構化整理，可進一步人工確認細節。")
        summary_lines.append("建議針對條列重點與數據欄位進行二次校對。")
    else:
        summary_lines.append("原始內容過短，無法形成完整摘要。")

    kps = "\n".join([f"- {p}" for p in key_points]) if key_points else "- （無可擷取重點）"
    datas = "\n".join([f"- {d}" for d in top_data]) if top_data else "- （未檢出明確數據）"

    return (
        "一、文件主題：\n"
        f"{topic}\n\n"
        "二、核心摘要：\n"
        + "\n".join(summary_lines[:5])
        + "\n\n三、關鍵重點：\n"
        + kps
        + "\n\n四、重要數據 / 結論：\n"
        + datas
        + "\n\n五、可行動建議（若適用）：\n"
        "- 先確認摘要中的主題與關鍵點是否符合原文目的。\n"
        "- 對重要數據進行人工覆核後再對外使用。"
    )


@csrf_exempt
@require_node("pdf", api=True)
def api_summary(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)

    doc_text = ""
    if request.content_type and "application/json" in request.content_type.lower():
        try:
            body = json.loads((request.body or b"").decode("utf-8") or "{}")
            doc_text = str(body.get("text") or "").strip()
        except Exception:
            doc_text = ""
    else:
        doc_text = (request.POST.get("text") or "").strip()

    if not doc_text:
        return JsonResponse({"ok": False, "error": "請先提供文件內容"}, status=400)

    max_chars = int(os.environ.get("PDF_SUMMARY_MAX_CHARS", "12000") or 12000)
    if len(doc_text) > max_chars:
        doc_text = doc_text[:max_chars]

    try:
        llm = get_chat_model(temperature=0.2, timeout=120)
        summary = _llm_to_text(llm.invoke(_build_summary_prompt(doc_text)))
        if not summary:
            summary = _local_summary_fallback(doc_text)
            return JsonResponse({"ok": True, "summary": summary, "fallback": True}, status=200)
        return JsonResponse({"ok": True, "summary": summary, "fallback": False}, status=200)
    except Exception:
        summary = _local_summary_fallback(doc_text)
        return JsonResponse({"ok": True, "summary": summary, "fallback": True}, status=200)

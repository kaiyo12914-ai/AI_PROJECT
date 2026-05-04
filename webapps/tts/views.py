from __future__ import annotations

import os
import json
import uuid
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any, Tuple, Optional

from django.conf import settings
from django.http import JsonResponse, HttpRequest, FileResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

from opencc import OpenCC

from webapps.portal.decorators import require_node


_cc = OpenCC("s2t")  # Simplified → Traditional


# =========================
# 頁面：/tts/
# =========================
@require_node("tts")
def index(request: HttpRequest):
    return render(request, "tts/index.html")


# =========================
# Piper 設定（可在 settings.py 覆寫）
# =========================
BASE_DIR = Path(__file__).resolve().parent.parent

PIPER_EXE = BASE_DIR / "piper" / "piper.exe"
PIPER_MODEL_DIR = BASE_DIR / "piper" / "models"
PIPER_OUTPUT_DIR = Path(settings.MEDIA_ROOT) / "tts"  # ⭐ 更標準

DEFAULT_MODEL_FILE = "zh_CN-huayan-medium.onnx"
ALLOWED_MODEL_SUFFIX = (".onnx",)


# -------------------------
# Export helpers (STT text -> files)
# -------------------------
@csrf_exempt
@require_node("tts", api=True)
@require_POST
def api_export_docx(request: HttpRequest):
    from docx import Document

    try:
        body = json.loads(request.body.decode("utf-8"))
    except Exception:
        return JsonResponse({"ok": False, "error": "invalid json"}, status=400)

    text = (body.get("text") or "").strip()
    if not text:
        return JsonResponse({"ok": False, "error": "text is required"}, status=400)

    out_dir = Path(settings.MEDIA_ROOT) / "stt"
    out_dir.mkdir(parents=True, exist_ok=True)

    file_id = uuid.uuid4().hex[:12]
    p = out_dir / f"stt_{file_id}.docx"

    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(p))

    return FileResponse(
        open(p, "rb"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@csrf_exempt
@require_node("tts", api=True)
@require_POST
def api_export_txt(request: HttpRequest):
    try:
        body = json.loads(request.body.decode("utf-8"))
    except Exception:
        return JsonResponse({"ok": False, "error": "invalid json"}, status=400)

    text = (body.get("text") or "").strip()
    if not text:
        return JsonResponse({"ok": False, "error": "text is required"}, status=400)

    out_dir = Path(settings.MEDIA_ROOT) / "stt"
    out_dir.mkdir(parents=True, exist_ok=True)

    file_id = uuid.uuid4().hex[:12]
    p = out_dir / f"stt_{file_id}.txt"
    p.write_text(text, encoding="utf-8")

    return FileResponse(open(p, "rb"), content_type="text/plain")


# =========================
# STT: whisper.cpp
# =========================
@csrf_exempt
@require_node("tts", api=True)
@require_POST
def api_stt_transcribe(request: HttpRequest):
    """
    POST multipart/form-data
    file: audio (wav/mp3/m4a...)
    return: {ok, text}
    """
    f = request.FILES.get("audio")
    if not f:
        return JsonResponse({"ok": False, "error": "audio file is required"}, status=400)
    language = (request.POST.get("language") or "auto").strip().lower()
    if language not in {"auto", "en", "zh"}:
        return JsonResponse({"ok": False, "error": "language must be auto, en, or zh"}, status=400)

    whisper_exe = Path(getattr(settings, "WHISPER_EXE", "")).resolve()
    whisper_model = Path(getattr(settings, "WHISPER_MODEL", "")).resolve()
    out_dir = Path(getattr(settings, "STT_OUTPUT_DIR", Path(settings.MEDIA_ROOT) / "stt")).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    if not whisper_exe.exists():
        return JsonResponse({"ok": False, "error": f"WHISPER_EXE not found: {whisper_exe}"}, status=500)
    if not whisper_model.exists():
        return JsonResponse({"ok": False, "error": f"WHISPER_MODEL not found: {whisper_model}"}, status=500)

    suffix = Path(f.name).suffix or ".wav"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        for chunk in f.chunks():
            tmp.write(chunk)
        tmp_path = Path(tmp.name)

    try:
        file_id = uuid.uuid4().hex[:12]
        out_prefix = out_dir / f"stt_{file_id}"

        cmd = [
            str(whisper_exe),
            "-m", str(whisper_model),
            "-f", str(tmp_path),
            "-of", str(out_prefix),
        ]
        if language in {"en", "zh"}:
            cmd += ["-l", language]
        cmd += ["-otxt"]

        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )

        if proc.returncode != 0:
            err = (proc.stderr or b"").decode("utf-8", errors="ignore").strip()
            out = (proc.stdout or b"").decode("utf-8", errors="ignore").strip()
            return JsonResponse(
                {"ok": False, "error": err or out or f"return code {proc.returncode}"},
                status=200,
            )

        txt_path = Path(str(out_prefix) + ".txt")
        if not txt_path.exists():
            return JsonResponse({"ok": False, "error": "transcribe ok but txt not found"}, status=200)

        text = txt_path.read_text(encoding="utf-8", errors="ignore").strip()

        try:
            text = _cc.convert(text)
        except Exception:
            pass

        return JsonResponse({"ok": True, "text": text, "language": language, "txt_path": str(txt_path)}, status=200)

    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


# =========================
# Piper internal helpers
# =========================
def _get_paths() -> Tuple[Path, Path, Path]:
    piper_exe = Path(getattr(settings, "PIPER_EXE", PIPER_EXE))
    model_dir = Path(getattr(settings, "PIPER_MODEL_DIR", PIPER_MODEL_DIR))
    out_dir = Path(getattr(settings, "PIPER_OUTPUT_DIR", PIPER_OUTPUT_DIR))
    return piper_exe, model_dir, out_dir


def _ensure_dirs(out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)


def _safe_model_path(model_dir: Path, model_name: str) -> Path:
    model_name = (model_name or "").strip()
    if not model_name:
        model_name = DEFAULT_MODEL_FILE

    if not model_name.lower().endswith(ALLOWED_MODEL_SUFFIX):
        raise ValueError("model 必須是 .onnx")

    model_name = os.path.basename(model_name)
    model_path = (model_dir / model_name).resolve()

    if model_dir.resolve() not in model_path.parents and model_path != model_dir.resolve():
        raise ValueError("model 路徑不合法")

    return model_path


def _read_payload(request: HttpRequest) -> Dict[str, Any]:
    if request.content_type and "application/json" in request.content_type.lower():
        try:
            return json.loads(request.body.decode("utf-8"))
        except Exception:
            return {}
    return {
        "text": request.POST.get("text", ""),
        "model": request.POST.get("model", ""),
        "speaker": request.POST.get("speaker", ""),
    }


def _run_piper(piper_exe: Path, model_path: Path, output_wav: Path, text: str) -> Tuple[bool, str]:
    if not piper_exe.exists():
        return False, f"piper.exe 不存在：{piper_exe}"
    if not model_path.exists():
        return False, f"模型檔不存在：{model_path}"

    cmd = [str(piper_exe), "--model", str(model_path), "--output_file", str(output_wav)]

    try:
        proc = subprocess.run(
            cmd,
            input=text.encode("utf-8"),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except Exception as e:
        return False, f"執行 piper 失敗：{e}"

    if proc.returncode != 0:
        err = (proc.stderr or b"").decode("utf-8", errors="ignore").strip()
        out = (proc.stdout or b"").decode("utf-8", errors="ignore").strip()
        return False, (err or out or f"piper return code {proc.returncode}")

    if not output_wav.exists() or output_wav.stat().st_size == 0:
        return False, "已執行但未產生 wav（檔案不存在或為空）"

    return True, "ok"


def _build_wav_url(out_wav: Path) -> Optional[str]:
    try:
        if settings.MEDIA_ROOT and settings.MEDIA_URL:
            media_root_path = Path(settings.MEDIA_ROOT).resolve()
            out_wav_resolve = out_wav.resolve()
            if media_root_path in out_wav_resolve.parents:
                rel = out_wav_resolve.relative_to(media_root_path).as_posix()
                return (settings.MEDIA_URL.rstrip("/") + "/" + rel).replace("//", "/")
    except Exception:
        pass
    return None


# =========================
# API：/tts/generate/
# POST {text, model?}
# =========================
@csrf_exempt
@require_node("tts", api=True)
@require_POST
def api_tts_generate(request: HttpRequest):
    payload = _read_payload(request)
    text = (payload.get("text") or "").strip()
    model_name = (payload.get("model") or "").strip()

    if not text:
        return JsonResponse({"ok": False, "error": "text is required"}, status=400)

    max_chars = int(getattr(settings, "TTS_MAX_CHARS", 1200))
    if len(text) > max_chars:
        return JsonResponse({"ok": False, "error": f"text too long (>{max_chars})"}, status=400)

    piper_exe, model_dir, out_dir = _get_paths()
    _ensure_dirs(out_dir)

    try:
        model_path = _safe_model_path(model_dir, model_name)
    except Exception as e:
        return JsonResponse({"ok": False, "error": f"invalid model: {e}"}, status=400)

    file_id = uuid.uuid4().hex[:12]
    out_wav = out_dir / f"tts_{file_id}.wav"

    ok, msg = _run_piper(piper_exe=piper_exe, model_path=model_path, output_wav=out_wav, text=text)

    if not ok:
        return JsonResponse(
            {"ok": False, "error": msg, "piper_exe": str(piper_exe), "model": str(model_path)},
            status=200,
        )

    return JsonResponse(
        {
            "ok": True,
            "model": model_path.name,
            "wav_path": str(out_wav),
            "wav_url": _build_wav_url(out_wav),
            "bytes": out_wav.stat().st_size,
            "message": "generated",
        },
        status=200,
    )


# =========================
# File -> Text extractors
# =========================
def _extract_text_txt(file_obj) -> str:
    data = file_obj.read()
    try:
        return data.decode("utf-8")
    except Exception:
        try:
            return data.decode("cp950", errors="ignore")
        except Exception:
            return data.decode("utf-8", errors="ignore")


def _extract_text_docx(file_obj) -> str:
    from docx import Document
    doc = Document(file_obj)
    lines = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines)


def _extract_text_pdf(file_obj) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file_obj)
    parts = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_text_by_ext(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    ext = os.path.splitext(name)[1]

    if ext in (".txt", ".csv"):
        return _extract_text_txt(uploaded_file)
    if ext == ".docx":
        return _extract_text_docx(uploaded_file)
    if ext == ".pdf":
        return _extract_text_pdf(uploaded_file)

    raise ValueError(f"unsupported file type: {ext}")


# =========================
# API：/tts/generate_from_file/
# POST multipart/form-data: file, model?
# =========================
@csrf_exempt
@require_node("tts", api=True)
@require_POST
def api_tts_generate_from_file(request: HttpRequest):
    f = request.FILES.get("file")
    if not f:
        return JsonResponse({"ok": False, "error": "file is required"}, status=400)

    model_name = (request.POST.get("model") or "").strip()

    max_mb = int(getattr(settings, "TTS_FILE_MAX_MB", 15))
    if f.size > max_mb * 1024 * 1024:
        return JsonResponse({"ok": False, "error": f"file too large (>{max_mb}MB)"}, status=400)

    try:
        text = (_extract_text_by_ext(f) or "").strip()
    except Exception as e:
        return JsonResponse({"ok": False, "error": f"extract failed: {e}"}, status=400)

    if not text:
        return JsonResponse(
            {
                "ok": False,
                "error": "file parsed but text is empty (PDF may be scanned image)",
                "filename": f.name,
            },
            status=200,
        )

    max_chars = int(getattr(settings, "TTS_MAX_CHARS", 1200))
    if len(text) > max_chars:
        text = text[:max_chars]

    piper_exe, model_dir, out_dir = _get_paths()
    _ensure_dirs(out_dir)

    try:
        model_path = _safe_model_path(model_dir, model_name)
    except Exception as e:
        return JsonResponse({"ok": False, "error": f"invalid model: {e}"}, status=400)

    file_id = uuid.uuid4().hex[:12]
    out_wav = out_dir / f"tts_{file_id}.wav"

    ok, msg = _run_piper(piper_exe=piper_exe, model_path=model_path, output_wav=out_wav, text=text)
    if not ok:
        return JsonResponse(
            {"ok": False, "error": msg, "piper_exe": str(piper_exe), "model": str(model_path)},
            status=200,
        )

    return JsonResponse(
        {
            "ok": True,
            "filename": f.name,
            "text": text,
            "model": model_path.name,
            "wav_path": str(out_wav),
            "wav_url": _build_wav_url(out_wav),
            "bytes": out_wav.stat().st_size,
            "message": "generated_from_file",
        },
        status=200,
    )

from __future__ import annotations

import csv
import io
import json
import os
from typing import Any, Dict, List, Tuple

import requests
from django.conf import settings
from django.http import HttpRequest, HttpResponse, JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

from webapps.doc.models import DocumentTemplate
from webapps.portal.decorators import require_node


# ============================================================
# Page
# ============================================================
@require_node("student")
def index(request: HttpRequest) -> HttpResponse:
    return render(request, "student/index.html")


# ============================================================
# Helpers
# ============================================================
def _safe_str(x: Any) -> str:
    return "" if x is None else str(x)


def _json_body(request: HttpRequest) -> Dict[str, Any]:
    try:
        raw = (request.body or b"").decode("utf-8").strip()
        return json.loads(raw) if raw else {}
    except Exception:
        return {}


def _download_items(request: HttpRequest) -> List[Tuple[str, str]]:
    """
    JSON:
    { "items": [ {"name":"...", "comment":"..."}, ... ] }
    """
    payload = _json_body(request)
    items = payload.get("items")
    if not isinstance(items, list):
        return []

    out: List[Tuple[str, str]] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        name = _safe_str(it.get("name")).strip()
        comment = _safe_str(it.get("comment")).strip()
        if name and comment:
            out.append((name, comment))
    return out


# ============================================================
# TTS API
# POST /student/tts/
# ============================================================
@csrf_exempt
@require_POST
@require_node("student", api=True)
def generate_comment_tts(request: HttpRequest) -> JsonResponse:
    draft = (
        request.POST.get("draft")
        or request.POST.get("text")
        or request.GET.get("draft")
        or request.GET.get("text")
        or ""
    ).strip()

    if not draft:
        return JsonResponse({"ok": False, "error": "draft/text is required"}, status=400)

    base = (getattr(settings, "TTS_API_BASE_URL", "") or "").rstrip("/")
    url = (base + "/tts/generate/") if base else "/tts/generate/"
    timeout = int(getattr(settings, "TTS_API_TIMEOUT", 60))

    try:
        r = requests.post(url, json={"text": draft}, timeout=timeout)
        data = r.json() if r.content else {}
    except Exception as e:
        return JsonResponse({"ok": False, "error": f"call tts api failed: {e}"}, status=500)

    if not data or not data.get("ok"):
        return JsonResponse(
            {"ok": False, "error": data.get("error", "tts failed"), "raw": data},
            status=500,
        )

    return JsonResponse(
        {"ok": True, "text": draft, "audio": data.get("wav_url"), "provider": "tts_api"},
        status=200,
    )


# ============================================================
# Attachment parse (DOCX / ODT / PDF)
# POST /student/parse/  multipart/form-data (attachments multiple)
# ============================================================
def _extract_text_by_ext(file_obj, filename: str) -> str:
    ext = os.path.splitext((filename or "").lower())[1]

    if ext == ".docx":
        from docx import Document

        doc = Document(file_obj)
        lines = [(p.text or "").strip() for p in doc.paragraphs]
        return "\n".join([x for x in lines if x])

    if ext == ".odt":
        from odf.opendocument import load
        from odf import text as odf_text

        odt = load(file_obj)
        lines: List[str] = []
        for p in odt.getElementsByType(odf_text.P):
            t = "".join([n.data for n in p.childNodes if getattr(n, "data", None)])
            t = (t or "").strip()
            if t:
                lines.append(t)
        return "\n".join(lines)

    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(file_obj)
        texts = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join([t for t in texts if t])

    raise ValueError(f"unsupported file type: {ext}")


@csrf_exempt
@require_node("student", api=True)
def api_parse_attachments(request: HttpRequest) -> JsonResponse:
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)

    files = request.FILES.getlist("attachments")
    if not files:
        return JsonResponse({"ok": False, "error": "no files uploaded"}, status=400)

    max_files = 5
    max_each_mb = 10

    results: List[Dict[str, Any]] = []
    combined_parts: List[str] = []

    for f in files[:max_files]:
        if f.size > max_each_mb * 1024 * 1024:
            results.append({"filename": f.name, "ok": False, "error": f"file too large (>{max_each_mb}MB)"})
            continue

        try:
            text = (_extract_text_by_ext(f, f.name) or "").strip()
            results.append({"filename": f.name, "ok": True, "chars": len(text), "text": text})
            if text:
                combined_parts.append(f"【附件：{f.name}】\n{text}")
        except Exception as e:
            results.append({"filename": f.name, "ok": False, "error": str(e)})

    return JsonResponse(
        {"ok": True, "files": results, "combined_text": "\n\n".join(combined_parts).strip()},
        status=200,
    )


# ============================================================
# Templates API (GET/POST)
# /student/templates/
# ============================================================
@csrf_exempt
@require_node("student", api=True)
def api_templates(request: HttpRequest) -> JsonResponse:
    if request.method == "GET":
        doc_type = (request.GET.get("doc_type") or "").strip()
        tag = (request.GET.get("tag") or "").strip()

        qs = DocumentTemplate.objects.all()
        if doc_type:
            qs = qs.filter(doc_type=doc_type)
        if tag:
            qs = qs.filter(tags__contains=[tag])

        data = [
            {
                "id": t.id,
                "title": t.title,
                "doc_type": t.doc_type,
                "description": t.description,
                "tags": t.tags or [],
                "content": t.content,
                "created_at": t.created_at.isoformat(),
                "updated_at": t.updated_at.isoformat(),
            }
            for t in qs.order_by("-created_at")
        ]
        return JsonResponse({"ok": True, "templates": data}, status=200)

    if request.method == "POST":
        body = _json_body(request)

        title = _safe_str(body.get("title")).strip()
        doc_type = _safe_str(body.get("doc_type")).strip()
        content = _safe_str(body.get("content")).strip()
        description = _safe_str(body.get("description")).strip()

        tags = body.get("tags") or []
        if not isinstance(tags, list):
            return JsonResponse({"ok": False, "error": "tags must be a list"}, status=400)

        # unique + strip
        cleaned: List[str] = []
        seen = set()
        for x in tags:
            s = _safe_str(x).strip()
            if s and s not in seen:
                seen.add(s)
                cleaned.append(s)

        if not title:
            return JsonResponse({"ok": False, "error": "missing field: title"}, status=400)
        if not doc_type:
            return JsonResponse({"ok": False, "error": "missing field: doc_type"}, status=400)
        if doc_type not in dict(DocumentTemplate.DOC_TYPE_CHOICES):
            return JsonResponse({"ok": False, "error": f"invalid doc_type: {doc_type}"}, status=400)
        if not content:
            return JsonResponse({"ok": False, "error": "missing field: content"}, status=400)

        t = DocumentTemplate.objects.create(
            title=title,
            doc_type=doc_type,
            description=description,
            tags=cleaned,
            content=content,
        )
        return JsonResponse({"ok": True, "message": "created", "id": t.id}, status=201)

    return JsonResponse({"ok": False, "error": "method not allowed"}, status=405)


# ============================================================
# Downloads: TXT / CSV / WORD(DOCX)
# ============================================================
@csrf_exempt
@require_POST
@require_node("student", api=True)
def download_txt(request: HttpRequest) -> HttpResponse:
    items = _download_items(request)

    buf = io.StringIO()
    for i, (name, comment) in enumerate(items, start=1):
        buf.write(name + "\n")
        buf.write(comment + "\n")
        if i != len(items):
            buf.write("\n")

    resp = HttpResponse(buf.getvalue(), content_type="text/plain; charset=utf-8")
    resp["Content-Disposition"] = 'attachment; filename="comments.txt"'
    return resp


@csrf_exempt
@require_POST
@require_node("student", api=True)
def download_csv(request: HttpRequest) -> HttpResponse:
    items = _download_items(request)

    out = io.StringIO()
    w = csv.writer(out)
    w.writerow(["name", "comment"])
    for name, comment in items:
        w.writerow([name, comment])

    resp = HttpResponse(out.getvalue(), content_type="text/csv; charset=utf-8")
    resp["Content-Disposition"] = 'attachment; filename="comments.csv"'
    return resp


@csrf_exempt
@require_POST
@require_node("student", api=True)
def download_word(request: HttpRequest) -> HttpResponse:
    from docx import Document  # python-docx

    items = _download_items(request)

    doc = Document()
    doc.add_heading("學生期末評語", level=1)

    for idx, (name, comment) in enumerate(items, start=1):
        doc.add_heading(f"{idx}. {name}", level=2)
        for line in (comment or "").splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph("")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    resp = HttpResponse(
        bio.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = 'attachment; filename="comments.docx"'
    return resp
